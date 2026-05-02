import os
from pathlib import Path
from typing import Tuple, Union

import pdfplumber
import docx
from bs4 import BeautifulSoup

# Optional OCR imports
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from app.config import PROCESSED_DIR, POPPLER_PATH
from app.utils.text_utils import clean_text


class DocumentProcessor:
    """Service for extracting text from different document types."""

    @staticmethod
    def extract_text_from_pdf(file_path: Path) -> str:
        """Extract text from PDF files with OCR fallback for image-based PDFs."""
        text = ""
        
        try:
            # First, try standard text extraction
            with pdfplumber.open(file_path) as pdf:
                if len(pdf.pages) == 0:
                    raise Exception("PDF file appears to be empty or corrupted")
                
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text() or ""
                        if page_text.strip():  # Only add non-empty pages
                            text += f"--- Page {page_num + 1} ---\n"
                            text += page_text + "\n\n"
                    except Exception as page_error:
                        print(f"Warning: Could not extract text from page {page_num + 1}: {page_error}")
                        continue
                
                # If we got some text, return it
                if text.strip():
                    return clean_text(text)
                
                # If no text extracted, try OCR if available
                if OCR_AVAILABLE:
                    print("No text found with standard extraction, trying OCR...")
                    return DocumentProcessor._extract_text_with_ocr(file_path)
                else:
                    # Return a minimal response instead of failing completely
                    return f"[PDF Document - {len(pdf.pages)} pages]\nNote: This appears to be an image-based PDF. Install OCR libraries (pytesseract, pdf2image) for text extraction."
                    
        except Exception as e:
            if OCR_AVAILABLE:
                try:
                    print(f"Standard PDF extraction failed: {e}. Trying OCR...")
                    return DocumentProcessor._extract_text_with_ocr(file_path)
                except Exception as ocr_error:
                    raise Exception(f"Both standard and OCR extraction failed. Standard: {str(e)}, OCR: {str(ocr_error)}")
            else:
                raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    @staticmethod
    def _extract_text_with_ocr(file_path: Path) -> str:
        """Extract text using OCR for image-based PDFs."""
        if not OCR_AVAILABLE:
            raise Exception("OCR libraries not available")
        
        try:
            # Convert PDF pages to images
            pages = convert_from_path(
            str(file_path),
            dpi=200,
            poppler_path=POPPLER_PATH
        )
            text_parts = []
            
            for page_num, page_image in enumerate(pages):
                try:
                    # Extract text using OCR
                    page_text = pytesseract.image_to_string(page_image, lang='eng')
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} (OCR) ---\n{page_text}\n")
                except Exception as ocr_error:
                    print(f"OCR failed for page {page_num + 1}: {ocr_error}")
                    continue
            
            if text_parts:
                return clean_text("\n".join(text_parts))
            else:
                return "[PDF Document - OCR extraction completed but no readable text found]"
                
        except Exception as e:
            raise Exception(f"OCR extraction failed: {str(e)}")
        
    @staticmethod
    def extract_text_from_docx(file_path: Path) -> str:
        """Extract text from DOCX files with enhanced content extraction."""
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                # Return a minimal response instead of failing
                return "[Word Document - No readable text content found]"
            
            full_text = "\n".join(text_parts)
            return clean_text(full_text)
            
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_path: Path) -> str:
        """Extract text from plain text files."""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                if not content.strip():
                    return "[Text file is empty]"
                return clean_text(content)
        except UnicodeDecodeError:
            try:
                # Try with a different encoding if UTF-8 fails
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                    if not content.strip():
                        return "[Text file is empty]"
                    return clean_text(content)
            except Exception as e:
                raise Exception(f"Error reading text file with multiple encodings: {str(e)}")
        except Exception as e:
            raise Exception(f"Error extracting text from text file: {str(e)}")
    
    @staticmethod
    def extract_text_from_html(file_path: Path) -> str:
        """Extract text from HTML files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.extract()
                
                text = soup.get_text()
                if not text.strip():
                    return "[HTML file contains no readable text]"
                    
                return clean_text(text)
        except Exception as e:
            raise Exception(f"Error extracting text from HTML: {str(e)}")
    
    @staticmethod
    def extract_text_from_markdown(file_path: Path) -> str:
        """Extract text from Markdown files (treated as plain text)."""
        return DocumentProcessor.extract_text_from_txt(file_path)
    
    @classmethod
    def process_document(
        cls, file_path: Union[str, Path]
    ) -> Tuple[str, int]:
        # Normalize to Path
        if isinstance(file_path, str):
            file_path = Path(file_path)

        # Existence & size checks
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        if file_path.stat().st_size > 50 * 1024 * 1024:
            raise ValueError("File too large (>50 MB)")

        ext = file_path.suffix.lower()
        # Unsupported extension
        if ext not in {".pdf", ".docx", ".txt", ".html", ".md"}:
            raise ValueError(f"Unsupported file format: {ext}")

        # Extract
        try:
            if ext == ".pdf":
                text = cls.extract_text_from_pdf(file_path)
            elif ext == ".docx":
                text = cls.extract_text_from_docx(file_path)
            elif ext == ".txt":
                text = cls.extract_text_from_txt(file_path)
            elif ext == ".html":
                text = cls.extract_text_from_html(file_path)
            else:
                text = cls.extract_text_from_markdown(file_path)
        except Exception:
            # On any extraction error, produce empty text
            text = ""

        # Char count and saving
        char_count = len(text)
        PROCESSED_DIR.mkdir(parents=True, exist_ok=True)
        out_name = f"{file_path.stem}_{ext[1:]}.txt"
        save_path = PROCESSED_DIR / out_name
        save_path.write_text(text, encoding="utf-8")

        return text, char_count

    @staticmethod
    def is_valid_file(filename: str) -> bool:
        return Path(filename).suffix.lower() in {
            ".pdf", ".docx", ".txt", ".html", ".md"
        }
